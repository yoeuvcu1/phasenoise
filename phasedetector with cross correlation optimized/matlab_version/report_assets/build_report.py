#!/usr/bin/env python3
"""Build the final Turkish Cross-PSD phase-noise report.

The source DOCX is opened as the document/template base, but it is never
modified.  The old partial body is replaced in memory, validated MATLAB
R2025b result artifacts are inserted, and the finished report is written to
the fixed output path under ``matlab_version``.

This script intentionally fails before authoring when a required summary,
raw MAT file, plot, source figure, or the final N=1,000,000 iteration sweep is
missing or inconsistent.
"""

from __future__ import annotations

import csv
import hashlib
import math
import os
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Mapping, Sequence

# h5py is kept as a task-local dependency so the bundled workspace Python and
# the user's system Python remain unchanged.  The directory is ignored by Git
# and is only needed while validating MATLAB -v7.3/HDF5 raw results.
LOCAL_PYTHON_DEPS = Path(__file__).resolve().parent / ".python_deps"
if LOCAL_PYTHON_DEPS.is_dir():
    sys.path.insert(0, str(LOCAL_PYTHON_DEPS))

try:
    import h5py
    import numpy as np
    from PIL import Image
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - environment diagnostic
    raise SystemExit(
        "Eksik Python bağımlılığı: "
        f"{exc}. Gerekli paketler: python-docx, Pillow ve h5py."
    ) from exc


# ---------------------------------------------------------------------------
# Paths and immutable experiment contracts
# ---------------------------------------------------------------------------

SCRIPT_PATH = Path(__file__).resolve()
REPORT_ASSETS_DIR = SCRIPT_PATH.parent
MATLAB_DIR = REPORT_ASSETS_DIR.parent
OPTIMIZED_DIR = MATLAB_DIR.parent
PROJECT_ROOT = OPTIMIZED_DIR.parent

SOURCE_DOCX = PROJECT_ROOT / "İki Kanallı Cross.docx"
OUTPUT_DOCX = MATLAB_DIR / "Iki_Kanalli_Cross_PSD_Faz_Gurultusu_Raporu.docx"
RESULTS_DIR = MATLAB_DIR / "results"
NARRATIVE_MD = REPORT_ASSETS_DIR / "REPORT_NARRATIVE.md"

SOURCE_COMMIT = "0799f9f"
COMPARISON_STAMP = "20260821_195439719"
REPORT_DATE = "21 Ağustos 2026"
RUNTIME_LABEL = "MATLAB R2025b"

RS_FIGURES = {
    "2_1": REPORT_ASSETS_DIR / "figures" / "rs_fig_2_1.png",
    "2_4": REPORT_ASSETS_DIR / "figures" / "rs_fig_2_4.png",
    "2_8": REPORT_ASSETS_DIR / "figures" / "rs_fig_2_8.png",
}

EXPECTED_COLUMNS = (
    "run_file",
    "value",
    "mean_abs_error_db",
    "correction_factor",
    "elapsed_s",
)

COMPARISON_BASE_CONFIG = {
    "N": 1_000_000.0,
    "fs": 1_000_000.0,
    "A": 1.0,
    "f0": 200_000.0,
    "settling_samples": 0.0,
    "lpf_cutoff": 200_000.0,
    "lpf_order": 4.0,
    "phase_rms_dut": 0.05,
    "phase_rms_ref1": 0.05,
    "phase_rms_ref2": 0.05,
    "number_of_iterations": 100.0,
    "number_of_log_bins": 100.0,
}

LONG_ITERATION_BASE_CONFIG = {
    "N": 1_000_000.0,
    "fs": 1_000_000.0,
    "A": 1.0,
    "f0": 200_000.0,
    "settling_samples": 0.0,
    "lpf_cutoff": 100_000.0,
    "lpf_order": 4.0,
    "phase_rms_dut": 0.02,
    "phase_rms_ref1": 0.05,
    "phase_rms_ref2": 0.05,
    "number_of_iterations": 100.0,
    "number_of_log_bins": 100.0,
}

COMPARISON_SPECS = {
    "lpf_cutoff": {
        "values": [
            1_000.0,
            5_000.0,
            10_000.0,
            25_000.0,
            50_000.0,
            75_000.0,
            100_000.0,
            200_000.0,
            300_000.0,
        ],
        "fields": ("lpf_cutoff",),
        "plot": "lpf_cutoff_comparison.png",
    },
    "rms_dut": {
        "values": [0.01, 0.02, 0.05, 0.10, 0.20, 0.50],
        "fields": ("phase_rms_dut",),
        "plot": "rms_dut_comparison.png",
    },
    "rms_ref": {
        "values": [0.01, 0.02, 0.05, 0.10, 0.20, 0.50],
        "fields": ("phase_rms_ref1", "phase_rms_ref2"),
        "plot": "rms_ref_comparison.png",
    },
    "iterations": {
        "values": [1.0, 10.0, 100.0, 200.0, 500.0, 1_000.0],
        "fields": ("number_of_iterations",),
        "plot": "iterations_comparison.png",
    },
    "log_bins": {
        "values": [10.0, 25.0, 50.0, 80.0, 100.0, 200.0],
        "fields": ("number_of_log_bins",),
        "plot": "log_bins_comparison.png",
    },
}

# The final run requested for this report.  A prior 11-point result exists,
# but its value list belongs to an older profile and must not be selected.
EXPECTED_FINAL_ITERATIONS = [
    1.0,
    10.0,
    100.0,
    250.0,
    500.0,
    1_000.0,
    5_000.0,
    10_000.0,
    20_000.0,
]


# ---------------------------------------------------------------------------
# Report data validation
# ---------------------------------------------------------------------------


class ReportBuildError(RuntimeError):
    """Raised when a report input is absent, stale, or inconsistent."""


@dataclass(frozen=True)
class SweepRow:
    run_file: str
    value: float
    mae_db: float
    correction_factor: float
    elapsed_s: float


@dataclass(frozen=True)
class SweepData:
    name: str
    directory: Path
    rows: tuple[SweepRow, ...]
    plot_path: Path

    @property
    def total_elapsed_s(self) -> float:
        return sum(row.elapsed_s for row in self.rows)


def require_file(path: Path, purpose: str) -> Path:
    if not path.is_file():
        raise ReportBuildError(f"Eksik {purpose}: {path}")
    return path


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def nearly_equal(left: float, right: float, tolerance: float = 1e-9) -> bool:
    return math.isclose(left, right, rel_tol=tolerance, abs_tol=tolerance)


def h5_scalar(handle: h5py.File, dataset_path: str) -> float:
    if dataset_path not in handle:
        raise ReportBuildError(
            f"MAT dosyasında beklenen alan bulunamadı: {handle.filename}::{dataset_path}"
        )
    value = handle[dataset_path][()]
    try:
        scalar = value.reshape(-1)[0]
    except AttributeError:
        scalar = value
    result = float(scalar)
    if not math.isfinite(result):
        raise ReportBuildError(
            f"MAT dosyasında sonlu olmayan değer: {handle.filename}::{dataset_path}"
        )
    return result


def h5_vector(handle: h5py.File, dataset_path: str) -> np.ndarray:
    if dataset_path not in handle:
        raise ReportBuildError(
            f"MAT dosyasında beklenen alan bulunamadı: {handle.filename}::{dataset_path}"
        )
    values = np.asarray(handle[dataset_path][()], dtype=float).reshape(-1)
    if values.size < 2 or not np.all(np.isfinite(values)):
        raise ReportBuildError(
            f"MAT dosyasında geçersiz vektör: {handle.filename}::{dataset_path}"
        )
    return values


def compute_band_mae(raw_path: Path, max_frequency_hz: float) -> float:
    """Return an auxiliary MAE from the official common 200-point log grid."""

    require_file(raw_path, "bant MAE ham MATLAB sonucu")
    try:
        with h5py.File(raw_path, "r") as handle:
            f_cross = h5_vector(handle, "/current_results/cross/frequency_binned")
            l_cross = h5_vector(handle, "/current_results/cross/phase_noise_binned")
            f_dut = h5_vector(handle, "/current_results/dut_fft/frequency_binned")
            l_dut = h5_vector(handle, "/current_results/dut_fft/phase_noise_binned")
    except OSError as exc:
        raise ReportBuildError(f"Bant MAE için MAT dosyası açılamadı: {raw_path}") from exc

    if f_cross.size != l_cross.size or f_dut.size != l_dut.size:
        raise ReportBuildError(f"Bant MAE frekans/seviye boyutu uyuşmuyor: {raw_path}")
    if (
        np.any(f_cross <= 0)
        or np.any(f_dut <= 0)
        or np.any(np.diff(f_cross) <= 0)
        or np.any(np.diff(f_dut) <= 0)
    ):
        raise ReportBuildError(f"Bant MAE frekans ekseni artan/pozitif değil: {raw_path}")

    f_min = max(float(f_cross.min()), float(f_dut.min()))
    f_max = min(float(f_cross.max()), float(f_dut.max()))
    if f_min >= f_max:
        raise ReportBuildError(f"Bant MAE için ortak frekans aralığı yok: {raw_path}")
    f_common = np.logspace(np.log10(f_min), np.log10(f_max), 200)
    cross_interp = np.interp(np.log10(f_common), np.log10(f_cross), l_cross)
    dut_interp = np.interp(np.log10(f_common), np.log10(f_dut), l_dut)
    band_mask = f_common <= float(max_frequency_hz)
    if not np.any(band_mask):
        raise ReportBuildError(f"Bant MAE sınırı ortak frekans ızgarasının altında: {raw_path}")
    result = float(np.mean(np.abs(cross_interp[band_mask] - dut_interp[band_mask])))
    if not math.isfinite(result):
        raise ReportBuildError(f"Bant MAE sonlu değil: {raw_path}")
    return result


def validate_png(path: Path, purpose: str) -> tuple[int, int]:
    require_file(path, purpose)
    try:
        with Image.open(path) as image:
            image.verify()
        with Image.open(path) as image:
            width, height = image.size
    except Exception as exc:  # Pillow provides format-specific exception types
        raise ReportBuildError(f"Geçersiz PNG ({purpose}): {path}: {exc}") from exc
    if width < 600 or height < 350:
        raise ReportBuildError(
            f"Rapor için yetersiz çözünürlüklü PNG ({width}x{height}): {path}"
        )
    return width, height


def read_summary_csv(path: Path) -> tuple[SweepRow, ...]:
    require_file(path, "summary.csv")
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        reader = csv.DictReader(stream)
        if tuple(reader.fieldnames or ()) != EXPECTED_COLUMNS:
            raise ReportBuildError(
                f"Beklenmeyen CSV şeması: {path}. "
                f"Beklenen={EXPECTED_COLUMNS}, bulunan={reader.fieldnames}"
            )
        rows: list[SweepRow] = []
        for line_number, record in enumerate(reader, start=2):
            try:
                row = SweepRow(
                    run_file=(record["run_file"] or "").strip(),
                    value=float(record["value"]),
                    mae_db=float(record["mean_abs_error_db"]),
                    correction_factor=float(record["correction_factor"]),
                    elapsed_s=float(record["elapsed_s"]),
                )
            except (TypeError, ValueError) as exc:
                raise ReportBuildError(
                    f"CSV satırı sayısal olarak okunamadı: {path}:{line_number}"
                ) from exc
            if not row.run_file:
                raise ReportBuildError(f"Boş run_file: {path}:{line_number}")
            numeric_values = (
                row.value,
                row.mae_db,
                row.correction_factor,
                row.elapsed_s,
            )
            if not all(math.isfinite(value) for value in numeric_values):
                raise ReportBuildError(f"Sonlu olmayan metrik: {path}:{line_number}")
            if row.value <= 0 or row.mae_db < 0 or row.correction_factor <= 0:
                raise ReportBuildError(f"Geçersiz metrik aralığı: {path}:{line_number}")
            rows.append(row)
    if not rows:
        raise ReportBuildError(f"Boş summary.csv: {path}")
    return tuple(rows)


def assert_expected_values(
    rows: Sequence[SweepRow], expected_values: Sequence[float], context: str
) -> None:
    actual_values = [row.value for row in rows]
    if len(actual_values) != len(expected_values) or any(
        not nearly_equal(actual, expected)
        for actual, expected in zip(actual_values, expected_values)
    ):
        raise ReportBuildError(
            f"{context} değer listesi güncel profil ile uyuşmuyor. "
            f"Beklenen={list(expected_values)}, bulunan={actual_values}"
        )
    if len(set(actual_values)) != len(actual_values):
        raise ReportBuildError(f"{context} içinde yinelenen test değeri var.")


def validate_raw_run(
    raw_path: Path,
    csv_row: SweepRow,
    expected_config: Mapping[str, float],
) -> None:
    require_file(raw_path, "ham MATLAB sonucu")
    try:
        with h5py.File(raw_path, "r") as handle:
            raw_value = h5_scalar(handle, "/value")
            raw_mae = h5_scalar(handle, "/current_results/mean_absolute_error_fft_db")
            raw_correction = h5_scalar(handle, "/current_results/correction_factor")
            raw_elapsed = h5_scalar(handle, "/elapsed_seconds_current")
            if not nearly_equal(raw_value, csv_row.value, 1e-10):
                raise ReportBuildError(
                    f"CSV/MAT test değeri uyuşmuyor: {raw_path}"
                )
            # CSV writes six decimals for metrics and three for elapsed time.
            if not nearly_equal(raw_mae, csv_row.mae_db, 6e-7):
                raise ReportBuildError(f"CSV/MAT MAE uyuşmuyor: {raw_path}")
            if not nearly_equal(raw_correction, csv_row.correction_factor, 6e-7):
                raise ReportBuildError(
                    f"CSV/MAT correction factor uyuşmuyor: {raw_path}"
                )
            if not nearly_equal(raw_elapsed, csv_row.elapsed_s, 6e-4):
                raise ReportBuildError(f"CSV/MAT süre uyuşmuyor: {raw_path}")
            for field_name, expected_value in expected_config.items():
                actual_value = h5_scalar(
                    handle, f"/current_results/config/{field_name}"
                )
                if not nearly_equal(actual_value, expected_value, 1e-10):
                    raise ReportBuildError(
                        "Config uyuşmazlığı: "
                        f"{raw_path.name}::{field_name}; "
                        f"beklenen={expected_value}, bulunan={actual_value}"
                    )
    except OSError as exc:
        raise ReportBuildError(f"MATLAB v7.3/HDF5 dosyası açılamadı: {raw_path}") from exc


def load_and_validate_sweep(
    directory: Path,
    name: str,
    expected_values: Sequence[float],
    base_config: Mapping[str, float],
    swept_fields: Sequence[str],
    plot_filename: str,
) -> SweepData:
    if not directory.is_dir():
        raise ReportBuildError(f"Eksik sonuç klasörü: {directory}")
    require_file(directory / "summary.mat", "summary.mat")
    rows = read_summary_csv(directory / "summary.csv")
    assert_expected_values(rows, expected_values, directory.name)
    plot_path = directory / "plots" / plot_filename
    validate_png(plot_path, f"{name} karşılaştırma grafiği")

    for row in rows:
        expected_config = dict(base_config)
        for field_name in swept_fields:
            expected_config[field_name] = row.value
        validate_raw_run(directory / "raw" / row.run_file, row, expected_config)

    return SweepData(
        name=name,
        directory=directory,
        rows=tuple(rows),
        plot_path=plot_path,
    )


def discover_final_iteration_sweep() -> SweepData:
    completed_candidates = sorted(
        directory
        for directory in RESULTS_DIR.glob("*_iterations")
        if directory.is_dir() and (directory / "summary.csv").is_file()
    )
    if not completed_candidates:
        raise ReportBuildError(
            "summary.csv içeren tamamlanmış bir *_iterations klasörü bulunamadı. "
            "run_iterations tamamen bitmeden rapor oluşturulamaz."
        )

    latest = completed_candidates[-1]
    try:
        return load_and_validate_sweep(
            directory=latest,
            name="final_iterations",
            expected_values=EXPECTED_FINAL_ITERATIONS,
            base_config=LONG_ITERATION_BASE_CONFIG,
            swept_fields=("number_of_iterations",),
            plot_filename="iterations_comparison.png",
        )
    except ReportBuildError as exc:
        raise ReportBuildError(
            "En yeni summary.csv içeren iterations klasörü final N=1.000.000 "
            f"profili değil veya henüz tamamlanmamış: {latest}. "
            "Beklenen değerler [1, 10, 100, 250, 500, 1000, 5000, "
            f"10000, 20000]. Ayrıntı: {exc}"
        ) from exc


def load_report_inputs() -> tuple[dict[str, SweepData], SweepData]:
    require_file(SOURCE_DOCX, "kaynak DOCX")
    require_file(NARRATIVE_MD, "güvenli anlatı notu")
    if not RESULTS_DIR.is_dir():
        raise ReportBuildError(f"Sonuç dizini bulunamadı: {RESULTS_DIR}")
    for figure_path in RS_FIGURES.values():
        validate_png(figure_path, "Rohde & Schwarz kaynak görseli")

    sweeps: dict[str, SweepData] = {}
    for sweep_name, spec in COMPARISON_SPECS.items():
        directory = RESULTS_DIR / f"{COMPARISON_STAMP}_{sweep_name}"
        sweeps[sweep_name] = load_and_validate_sweep(
            directory=directory,
            name=sweep_name,
            expected_values=spec["values"],
            base_config=COMPARISON_BASE_CONFIG,
            swept_fields=spec["fields"],
            plot_filename=spec["plot"],
        )

    final_iterations = discover_final_iteration_sweep()
    return sweeps, final_iterations


# ---------------------------------------------------------------------------
# Word theme and OOXML helpers
# ---------------------------------------------------------------------------


NAVY = "0B2748"
NAVY_2 = "173D63"
TURQUOISE = "00A6A6"
TURQUOISE_DARK = "007C83"
LIGHT_TURQUOISE = "E8F7F7"
LIGHT_BLUE = "EEF3F8"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "5B6573"
DARK_TEXT = "1D2733"
WHITE = "FFFFFF"
BORDER_GRAY = "CFD8E3"
FONT_FAMILY = "Arial"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_description(table, title: str, description: str) -> None:
    """Add semantic metadata to layout tables without inventing a header row."""
    table_pr = table._tbl.tblPr
    caption = table_pr.find(qn("w:tblCaption"))
    if caption is None:
        caption = OxmlElement("w:tblCaption")
        table_pr.append(caption)
    caption.set(qn("w:val"), title)

    table_description = table_pr.find(qn("w:tblDescription"))
    if table_description is None:
        table_description = OxmlElement("w:tblDescription")
        table_pr.append(table_description)
    table_description.set(qn("w:val"), description)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER_GRAY, size="6") -> None:
    table_pr = table._tbl.tblPr
    borders = table_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_paragraph_border(paragraph, side: str, color: str, size="10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = borders.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        borders.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)


def add_word_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instruction_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    if placeholder:
        paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def clear_container(container) -> None:
    element = container._element
    for child in list(element):
        element.remove(child)


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_style_font(style, name: str, size_pt: float, color: str, bold=False) -> None:
    style.font.name = name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "tr-TR")


def get_or_create_style(document: Document, name: str, style_type) -> object:
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_style_font(normal, FONT_FAMILY, 10.5, DARK_TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    title = get_or_create_style(document, "Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title, FONT_FAMILY, 25, NAVY, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(12)

    subtitle = get_or_create_style(document, "Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(subtitle, FONT_FAMILY, 13, TURQUOISE_DARK, bold=True)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)

    heading_1 = get_or_create_style(document, "Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(heading_1, FONT_FAMILY, 16.5, NAVY, bold=True)
    heading_1.paragraph_format.space_before = Pt(4)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = get_or_create_style(document, "Heading 2", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(heading_2, FONT_FAMILY, 12.5, TURQUOISE_DARK, bold=True)
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = get_or_create_style(document, "Heading 3", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(heading_3, FONT_FAMILY, 10.8, NAVY_2, bold=True)
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True

    caption = get_or_create_style(document, "Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(caption, FONT_FAMILY, 8.5, MID_GRAY)
    caption.font.italic = True
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True

    list_paragraph = get_or_create_style(document, "List Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(list_paragraph, FONT_FAMILY, 10.2, DARK_TEXT)
    list_paragraph.paragraph_format.space_after = Pt(3)

    list_number = get_or_create_style(document, "List Number", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(list_number, FONT_FAMILY, 10.2, DARK_TEXT)
    list_number.paragraph_format.left_indent = Cm(0.65)
    list_number.paragraph_format.first_line_indent = Cm(-0.35)
    list_number.paragraph_format.space_after = Pt(3)

    cover_label = get_or_create_style(document, "Cover Label", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(cover_label, FONT_FAMILY, 10.5, TURQUOISE_DARK, bold=True)
    cover_label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_label.paragraph_format.space_after = Pt(10)

    lead = get_or_create_style(document, "Lead", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(lead, FONT_FAMILY, 11.2, NAVY_2)
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead.paragraph_format.line_spacing = 1.15
    lead.paragraph_format.space_after = Pt(9)

    equation = get_or_create_style(document, "Equation", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(equation, "Cambria Math", 11.5, NAVY)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(5)
    equation.paragraph_format.space_after = Pt(7)
    equation.paragraph_format.keep_together = True

    toc_entry = get_or_create_style(document, "Manual TOC", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(toc_entry, FONT_FAMILY, 10.5, NAVY_2)
    toc_entry.paragraph_format.space_after = Pt(4)

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)


def configure_section(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def configure_header_footer(document: Document) -> None:
    section = document.sections[0]
    clear_container(section.header)
    header = section.header.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("TASNİF DIŞI")
    header_run.bold = True
    header_run.font.name = FONT_FAMILY
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor.from_string(NAVY)
    set_paragraph_border(header, "bottom", TURQUOISE, size="8")

    clear_container(section.footer)
    footer = section.footer.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    set_paragraph_border(footer, "top", TURQUOISE, size="8")
    lead = footer.add_run("TASNİF DIŞI   •   Faz Gürültüsü Analizi   •   Sayfa ")
    lead.font.name = FONT_FAMILY
    lead.font.size = Pt(8)
    lead.font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_word_field(footer, " PAGE ", "1")
    footer.add_run(" / ")
    add_word_field(footer, " NUMPAGES ", "1")
    for run in footer.runs:
        run.font.name = FONT_FAMILY
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_body_paragraph(document: Document, text: str, style=None):
    paragraph = document.add_paragraph(text, style=style)
    if style is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_bullet(document: Document, text: str, level: int = 0):
    marker = "•" if level == 0 else "◦"
    paragraph = document.add_paragraph(style="List Paragraph")
    paragraph.add_run(f"{marker} {text}")
    paragraph.paragraph_format.left_indent = Cm(0.65 + 0.45 * level)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_numbered(document: Document, number: int, text: str):
    paragraph = document.add_paragraph(style="List Paragraph")
    paragraph.add_run(f"{number}. {text}")
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def add_equation(document: Document, text: str):
    return document.add_paragraph(text, style="Equation")


def add_callout(document: Document, title: str, text: str, color=LIGHT_TURQUOISE):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.2)
    set_table_description(
        table,
        title,
        f"{title}: {text}",
    )
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)
    set_table_borders(table, TURQUOISE, size="10")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    title_run = paragraph.add_run(f"{title}: ")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(TURQUOISE_DARK)
    body_run = paragraph.add_run(text)
    body_run.font.color.rgb = RGBColor.from_string(DARK_TEXT)
    for run in paragraph.runs:
        run.font.name = FONT_FAMILY
        run.font.size = Pt(9.7)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def set_table_widths(table, widths_cm: Sequence[float]) -> None:
    table.autofit = False
    for index, width_cm in enumerate(widths_cm):
        width = Cm(width_cm)
        table.columns[index].width = width
        for cell in table.columns[index].cells:
            cell.width = width


def add_data_table(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_cm: Sequence[float],
    numeric_columns: Iterable[int] = (),
):
    if len(headers) != len(widths_cm):
        raise ReportBuildError("Tablo başlığı ve genişlik sayısı uyuşmuyor.")
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths_cm)
    set_table_borders(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    prevent_row_split(header_row)
    for index, header_text in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=120, bottom=120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header_text)
        run.bold = True
        run.font.name = FONT_FAMILY
        run.font.size = Pt(8.6)
        run.font.color.rgb = RGBColor.from_string(WHITE)

    numeric_set = set(numeric_columns)
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for column_index, value in enumerate(row_values):
            cell = row.cells[column_index]
            if row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if column_index in numeric_set
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(str(value))
            run.font.name = FONT_FAMILY
            run.font.size = Pt(8.4)
            run.font.color.rgb = RGBColor.from_string(DARK_TEXT)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(
    document: Document,
    image_path: Path,
    caption: str,
    max_width_cm: float = 16.2,
    max_height_cm: float = 13.0,
):
    validate_png(image_path, "rapor görseli")
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    ratio = width_px / height_px
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    picture = paragraph.add_run().add_picture(
        str(image_path), width=Cm(width_cm), height=Cm(height_cm)
    )
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(".", 1)[0])
    caption_paragraph = document.add_paragraph(caption, style="Caption")
    caption_paragraph.paragraph_format.keep_together = True
    return paragraph, caption_paragraph


def add_section_title(document: Document, title: str, page_break=True):
    if page_break:
        document.add_page_break()
    heading = document.add_heading(title, level=1)
    set_paragraph_border(heading, "bottom", TURQUOISE, size="10")
    return heading


def add_subheading(document: Document, title: str, level=2):
    return document.add_heading(title, level=level)


def add_process_chain(document: Document) -> None:
    labels = [
        "DUT / Ref\n1/f³ gürültü",
        "İki kanal\nmikser",
        "Butterworth\nLPF ve /Kpd",
        "FFT tabanlı\nkompleks Cross-PSD",
        "Yineleme ortalaması\nlog-bin ve dBc/Hz",
    ]
    table = document.add_table(rows=1, cols=len(labels))
    set_table_description(
        table,
        "MATLAB işlem zinciri",
        "DUT ve iki referans girişinden Cross-PSD, yineleme ortalaması ve logaritmik binleme çıktısına uzanan işlem zinciri.",
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_widths(table, [3.15, 3.15, 3.15, 3.15, 3.15])
    for index, label in enumerate(labels):
        cell = table.cell(0, index)
        set_cell_shading(cell, NAVY if index % 2 == 0 else TURQUOISE_DARK)
        set_cell_margins(cell, top=180, bottom=180)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        run.bold = True
        run.font.name = FONT_FAMILY
        run.font.size = Pt(8.4)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    set_table_borders(table, WHITE, size="12")
    caption = document.add_paragraph(
        "Şekil 4.1. MATLAB modelinin uçtan uca işlem zinciri. Kaynak: Bu çalışma.",
        style="Caption",
    )
    caption.paragraph_format.keep_together = True


# ---------------------------------------------------------------------------
# Formatting and analysis helpers
# ---------------------------------------------------------------------------


def fmt_decimal(value: float, digits: int = 3) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def fmt_int(value: float) -> str:
    return f"{int(round(value)):,}".replace(",", ".")


def fmt_duration(seconds: float) -> str:
    if seconds < 120:
        return f"{fmt_decimal(seconds, 3)} s"
    minutes = int(seconds // 60)
    remainder = seconds - minutes * 60
    return f"{minutes} dk {fmt_decimal(remainder, 1)} s"


def get_row(sweep: SweepData, value: float) -> SweepRow:
    for row in sweep.rows:
        if nearly_equal(row.value, value):
            return row
    raise ReportBuildError(f"{sweep.name} içinde değer bulunamadı: {value}")


def best_row(sweep: SweepData) -> SweepRow:
    return min(sweep.rows, key=lambda row: row.mae_db)


def worst_row(sweep: SweepData) -> SweepRow:
    return max(sweep.rows, key=lambda row: row.mae_db)


def increasing_steps(rows: Sequence[SweepRow]) -> list[tuple[float, float]]:
    return [
        (previous.value, current.value)
        for previous, current in zip(rows, rows[1:])
        if current.mae_db > previous.mae_db + 1e-12
    ]


def describe_steps(steps: Sequence[tuple[float, float]]) -> str:
    if not steps:
        return "MAE artışı görülmemiştir"

    def format_step_value(value: float) -> str:
        if abs(value) < 1:
            return fmt_decimal(value, 2)
        return fmt_int(value)

    formatted_steps = [
        f"{format_step_value(left)}→{format_step_value(right)}"
        for left, right in steps
    ]
    if len(formatted_steps) == 1:
        return f"MAE artışı {formatted_steps[0]} geçişinde görülmüştür"
    return (
        "MAE artışları "
        + ", ".join(formatted_steps)
        + " geçişlerinde görülmüştür"
    )


def result_table_rows(sweep: SweepData, value_formatter) -> list[list[str]]:
    return [
        [
            value_formatter(row.value),
            fmt_decimal(row.mae_db, 3),
            fmt_decimal(row.correction_factor, 6),
            fmt_decimal(row.elapsed_s, 3),
        ]
        for row in sweep.rows
    ]


def start_landscape_section(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=True)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    return section


def start_portrait_section(document: Document):
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, landscape=False)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    return section


# ---------------------------------------------------------------------------
# Report content
# ---------------------------------------------------------------------------


def build_cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(32)
    label = document.add_paragraph("TEKNİK PROJE RAPORU", style="Cover Label")
    set_paragraph_border(label, "bottom", TURQUOISE, size="14")

    title = document.add_paragraph(
        "İki Kanallı Cross-PSD Yöntemiyle Faz Gürültüsü Ölçümünün Simülasyonu",
        style="Title",
    )
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(15)

    subtitle = document.add_paragraph(
        "MATLAB Modeli, Algoritmik Optimizasyon ve Parametrik İnceleme",
        style="Subtitle",
    )
    set_paragraph_border(subtitle, "bottom", NAVY, size="6")

    document.add_paragraph().paragraph_format.space_after = Pt(32)
    add_data_table(
        document,
        ["Rapor bilgisi", "Doğrulanmış değer"],
        [
            ["Çalışma ortamı", RUNTIME_LABEL],
            ["Kaynak senkron commit'i", SOURCE_COMMIT],
            ["Rapor tarihi", REPORT_DATE],
        ],
        [6.0, 10.0],
    )

    notice = document.add_paragraph("TASNİF DIŞI")
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(48)
    run = notice.runs[0]
    run.bold = True
    run.font.name = FONT_FAMILY
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def build_front_matter(
    document: Document, sweeps: Mapping[str, SweepData], final_iterations: SweepData
) -> None:
    comparison_iterations = sweeps["iterations"]
    comparison_first = comparison_iterations.rows[0]
    comparison_last = comparison_iterations.rows[-1]
    final_first = final_iterations.rows[0]
    final_last = final_iterations.rows[-1]
    final_low_band_mae = compute_band_mae(
        final_iterations.directory / "raw" / final_last.run_file, 10_000.0
    )
    lpf_best = best_row(sweeps["lpf_cutoff"])
    ref_high = get_row(sweeps["rms_ref"], 0.50)

    add_section_title(document, "Özet", page_break=True)
    add_body_paragraph(
        document,
        "Bu çalışmada, bir test edilen cihazın (Device Under Test, DUT) faz "
        "gürültüsünü iki ayrı referans kanalı üzerinden kestiren iki kanallı "
        "çapraz korelasyon yöntemi MATLAB R2025b ortamında modellenmiştir. "
        "DUT iki kanalda ortak tutulmuş, referanslar ayrı rastgele dizilerden "
        "üretilmiş ve kompleks çapraz güç spektral yoğunlukları (Cross-PSD) "
        "büyüklük alınmadan önce yinelemeler boyunca ortalanmıştır.",
        style="Lead",
    )
    add_body_paragraph(
        document,
        "Benzetim; sentetik 1/f³ faz gürültüsü üretimi, iki mikser, dördüncü "
        "derece Butterworth alçak geçiren filtre, faz dedektörü kazanç "
        "normalizasyonu, FFT tabanlı Cross-PSD, doğrusal DUT periodogram "
        "ortalaması ve logaritmik binleme bloklarından oluşmaktadır. Açık "
        "xcorr→ifftshift→fft zinciri korelasyon teoremine dayalı doğrudan "
        "spektral çarpımla değiştirilmiş; FFT boyu 2^nextpow2(2M−1) olarak "
        "seçilmiş ve tekrarlı hesaplamalar sadeleştirilmiştir. Kontrollü bir "
        "benchmark bulunmadığı için sayısal hızlanma çarpanı raporlanmamıştır.",
    )
    add_body_paragraph(
        document,
        "N=1.000.000 örnekli genel karşılaştırma profilinde yineleme sayısı "
        f"{fmt_int(comparison_first.value)}'den {fmt_int(comparison_last.value)}'e "
        f"çıkarıldığında tam bant MAE {fmt_decimal(comparison_first.mae_db)} "
        f"dB'den {fmt_decimal(comparison_last.mae_db)} dB'ye değişmiştir. "
        "Farklı config kullanan özel uzun taramada aynı metrik "
        f"{fmt_decimal(final_first.mae_db)} dB'den "
        f"{fmt_decimal(final_last.mae_db)} dB'ye düşmüştür. LPF taramasında "
        f"en düşük gözlenen MAE {fmt_int(lpf_best.value / 1_000)} kHz'de "
        f"{fmt_decimal(lpf_best.mae_db)} dB iken, 0,50 rad referans RMS "
        f"koşulunda MAE {fmt_decimal(ref_high.mae_db)} dB olmuştur. "
        f"20.000 yineleme kaydında 10 kHz ve altıyla sınırlanan ek tanısal "
        f"MAE {fmt_decimal(final_low_band_mae)} dB'dir. Bu değerler tek "
        "koşulu Monte Carlo sonuçlarıdır; genel optimum, donanım "
        "doğruluğu veya güven aralığı olarak yorumlanmamıştır.",
    )
    keyword = document.add_paragraph()
    keyword.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lead = keyword.add_run("Anahtar kelimeler: ")
    lead.bold = True
    keyword.add_run(
        "faz gürültüsü, çapraz korelasyon, Cross-PSD, FFT, faz dedektörü, "
        "Monte Carlo, MATLAB, dBc/Hz"
    )

    add_section_title(document, "İçindekiler", page_break=True)
    contents = [
        ("1", "Giriş ve motivasyon"),
        ("2", "Faz gürültüsü temelleri"),
        ("3", "Faz gürültüsü ölçüm yöntemleri"),
        ("4", "İki kanallı MATLAB benzetim modeli"),
        ("5", "Projenin gelişimi ve algoritmik optimizasyonlar"),
        ("6", "Deney tasarımı ve sonuç sözleşmesi"),
        ("7", "N=1.000.000 benzetim sonuçları"),
        ("8", "Tartışma ve sınırlamalar"),
        ("9", "Sonuç ve gelecek çalışmalar"),
        ("", "Kaynakça"),
        ("Ek A", "Sonuç provenance ve config manifesti"),
        ("Ek B", "Yazılım bileşenleri ve sorumlulukları"),
    ]
    toc_table = document.add_table(rows=0, cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    toc_table.autofit = False
    set_table_description(
        toc_table,
        "İçindekiler",
        "Rapor bölüm numaralarını ve bölüm adlarını eşleyen içerik listesi.",
    )
    set_table_widths(toc_table, [2.1, 13.8])
    for number, title in contents:
        row = toc_table.add_row()
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, top=90, bottom=90)
        p_number = row.cells[0].paragraphs[0]
        p_number.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_number = p_number.add_run(number)
        run_number.bold = True
        run_number.font.color.rgb = RGBColor.from_string(TURQUOISE_DARK)
        p_title = row.cells[1].paragraphs[0]
        run_title = p_title.add_run(title)
        run_title.font.color.rgb = RGBColor.from_string(NAVY_2)
        for paragraph in (p_number, p_title):
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = FONT_FAMILY
                run.font.size = Pt(10.3)

    add_subheading(document, "Kısaltmalar", level=2)
    add_data_table(
        document,
        ["Kısaltma", "Açılım / anlam"],
        [
            ["DUT", "Device Under Test - test edilen cihaz"],
            ["PSD", "Power Spectral Density - güç spektral yoğunluğu"],
            ["SSB", "Single Sideband - tek yan bant"],
            ["LPF", "Low-Pass Filter - alçak geçiren filtre"],
            ["MAE", "Ortalama mutlak dB farkı"],
            ["FFT", "Fast Fourier Transform"],
            ["RNG", "Rastgele sayı üreteci"],
        ],
        [3.0, 13.0],
    )


def build_theory_and_methods(document: Document) -> None:
    add_section_title(document, "1. Giriş ve motivasyon", page_break=True)
    add_body_paragraph(
        document,
        "İdeal bir osilatör frekans bölgesinde tek bir spektral çizgi üretir. "
        "Gerçek devrelerde termal etkiler, aktif eleman gürültüsü, rezonatör "
        "kayıpları, besleme değişimleri ve çevresel koşullar sinyalin anlık "
        "fazında rastgele sapmalara yol açar. Taşıyıcı çevresinde oluşan bu "
        "yan bantlar faz gürültüsü olarak adlandırılır.",
        style="Lead",
    )
    add_body_paragraph(
        document,
        "Faz gürültüsü haberleşme sistemlerinde modülasyon kalitesi ve komşu "
        "kanal performansını, radar sistemlerinde güçlü yansımalar yakınındaki "
        "zayıf hedeflerin seçilebilirliğini, sayısal sistemlerde ise saat "
        "jitteri ve örnekleme doğruluğunu sınırlar. Bu nedenle düşük gürültülü "
        "bir osilatörü ölçebilmek, yalnız DUT performansına değil ölçüm "
        "sisteminin kendi gürültü tabanına da bağlıdır.",
    )
    add_subheading(document, "1.1. Amaç ve kapsam")
    add_body_paragraph(
        document,
        "Çalışmanın amacı iki referans kanallı faz dedektörü mimarisini "
        "sayısal olarak kurmak, kanallarda ortak DUT bileşenini Cross-PSD "
        "ortalamasıyla kestirmek ve temel parametrelerin uçtan uca spektral "
        "uyuma etkisini incelemektir. Model bir ticari analizörün bütün "
        "donanım ayrıntılarını değil, yöntemin istatistiksel ve algoritmik "
        "çekirdeğini temsil eder.",
    )
    add_callout(
        document,
        "Çalışmanın katkısı",
        "Aynı Monte Carlo popülasyonundan üretilen Cross-PSD ve filtresiz DUT "
        "periodogramını karşılaştıran, uzun taramaları ham MAT/CSV/PNG kanıt "
        "zinciriyle saklayan modüler bir MATLAB modeli oluşturulmuştur.",
    )

    add_section_title(document, "2. Faz gürültüsü temelleri", page_break=True)
    add_subheading(document, "2.1. Zaman ve frekans bölgesi tanımı")
    add_body_paragraph(
        document,
        "Yalnız faz sapmasının modellendiği bir taşıyıcı aşağıdaki biçimde "
        "ifade edilir. Burada A nominal genliği, f₀ taşıyıcı frekansını ve "
        "φ(t) zamana bağlı faz hatasını göstermektedir.",
    )
    add_equation(document, "x(t) = A cos(2πf₀t + φ(t))")
    add_figure(
        document,
        RS_FIGURES["2_1"],
        "Şekil 2.1. İdeal ve faz gürültülü işaretin zaman/frekans alanındaki "
        "karşılaştırması. Kaynak: Rohde & Schwarz, Mastering Phase Noise "
        "Measurements, Part 1, Fig. 2-1, birleşik PDF s. 4.",
        max_height_cm=9.5,
    )

    add_subheading(document, "2.2. Tek yan bant faz gürültüsü")
    add_body_paragraph(
        document,
        "Faz gürültüsü çoğunlukla taşıyıcıdan belirli bir offset frekansındaki "
        "1 Hz bant genişliğine normalize edilmiş tek yan bant güç oranı olarak "
        "verilir. Sφ(f) tek taraflı faz dalgalanması PSD'si olarak "
        "tanımlandığında, küçük faz yaklaşımındaki SSB bağıntısı şöyledir:",
    )
    add_equation(document, "L(f) = 10 log₁₀(Sφ(f) / 2)   [dBc/Hz]")
    add_figure(
        document,
        RS_FIGURES["2_4"],
        "Şekil 2.2. Offset frekansı ve dBc/Hz eksenleriyle faz gürültüsünün "
        "nicelendirilmesi. Kaynak: Rohde & Schwarz, Part 1, Fig. 2-4, "
        "birleşik PDF s. 8.",
        max_height_cm=9.0,
    )

    add_subheading(document, "2.3. Güç yasası bölgeleri")
    add_body_paragraph(
        document,
        "Gerçek osilatör spektrumları birden fazla fiziksel mekanizmanın "
        "bileşimidir. Yaklaşık güç yasası davranışları aşağıda özetlenmiştir. "
        "Bu projede yalnız flicker FM karakterini temsil eden 1/f³ faz PSD "
        "bileşeni kullanılmıştır.",
    )
    add_data_table(
        document,
        ["Gürültü türü", "PSD davranışı", "Yaklaşık eğim", "Baskın bölge"],
        [
            ["Random-walk FM", "1/f⁴", "−40 dB/dekad", "Taşıyıcıya çok yakın"],
            ["Flicker FM", "1/f³", "−30 dB/dekad", "Yakın offsetler"],
            ["White FM", "1/f²", "−20 dB/dekad", "Orta offsetler"],
            ["Flicker PM", "1/f", "−10 dB/dekad", "Orta/uzak offsetler"],
            ["White PM", "Sabit", "0 dB/dekad", "Uzak offsetler"],
        ],
        [4.1, 3.0, 3.4, 5.5],
        numeric_columns=(1, 2),
    )
    add_body_paragraph(
        document,
        "Üretici, beyaz Gauss gürültüsünün FFT genliğini 1/√(k³) ile "
        "şekillendirir, DC bileşenini kaldırır ve zaman dizisini hedef RMS "
        "değerine normalize eder. Bu sentetik model ayrık spur'ları veya "
        "farklı eğim bölgelerini birlikte üretmez.",
    )

    add_section_title(document, "3. Faz gürültüsü ölçüm yöntemleri", page_break=True)
    add_subheading(document, "3.1. Doğrudan spektral analiz")
    add_body_paragraph(
        document,
        "Doğrudan yöntemde DUT çıkışı spektrum analizörüne uygulanır ve "
        "taşıyıcı çevresindeki gürültü yan bantları ölçülür. Kurulumu basittir; "
        "ancak ölçülebilecek en düşük seviye analizörün yerel osilatörü ve iç "
        "gürültü tabanı tarafından sınırlandırılır.",
    )
    add_subheading(document, "3.2. Faz dedektörü yöntemi")
    add_body_paragraph(
        document,
        "DUT, aynı frekansta ve nominal olarak 90° merkez faz farkında çalışan "
        "bir referansla mikserde çarpılır. LPF sonrasında küçük faz farkına "
        "orantılı taban bant bileşeni elde edilir. Taşıyıcının bastırılması "
        "duyarlılığı artırabilir; ancak tek kanal çıkışı hem DUT hem referans "
        "gürültüsünü içerir.",
    )
    add_subheading(document, "3.3. İki kanallı çapraz korelasyon yöntemi")
    add_body_paragraph(
        document,
        "İki kanallı yaklaşım faz dedektörü zincirini ikinci bir referansla "
        "çoğaltır. Aynı DUT iki kanalda ortak, referanslara ait terimler ise "
        "modelde korelasyonsuz kabul edilir. Çoklu ölçüm ortalamasında ortak "
        "DUT bileşeni korunurken kanala özgü terimler azalır.",
    )
    add_figure(
        document,
        RS_FIGURES["2_8"],
        "Şekil 3.1. İki referanslı faz dedektörü ve çapraz korelasyon mimarisi. "
        "Kaynak: Rohde & Schwarz, Part 2, Fig. 2-8, birleşik PDF s. 20.",
        max_height_cm=10.0,
    )
    add_body_paragraph(
        document,
        "Küçük faz farkı bölgesinde iki ölçüm kanalı ortak bir işaret "
        "katsayısıyla yaklaşık olarak aşağıdaki biçimde yazılabilir:",
    )
    add_equation(
        document,
        "y₁(t)=s[φD(t)−φR1(t)],   y₂(t)=s[φD(t)−φR2(t)],   s∈{−1,+1}",
    )
    add_equation(document, "S₁₂(f)=E{Y₁(f)Y₂*(f)} ≈ SD(f)")
    add_body_paragraph(
        document,
        "Sonlu kayıtta çapraz terimler tam olarak sıfırlanmaz. İdeal bağımsız "
        "kanal varsayımında ortalama sayısı K arttıkça gürültü tabanı "
        "iyileşmesinin yaklaşık 5 log₁₀(K) dB mertebesinde olması beklenir. "
        "R&S örneğinde 100 korelasyon için 10 dB, 10.000 korelasyon için 20 "
        "dB referans gürültüsü azalması verilir. Bu bağıntı bu rapordaki geniş "
        "bant MAE metriğiyle aynı büyüklük değildir.",
    )
    add_subheading(document, "3.4. Diğer yöntemler")
    add_body_paragraph(
        document,
        "Gecikme hattı ayırıcısı frekans değişimlerini faz değişimine dönüştürür; "
        "sayısal faz demodülasyonu ise örneklenmiş I/Q bileşenlerinden anlık "
        "fazı hesaplar. Bu yöntemler farklı uygulamalarda avantajlı olsa da "
        "çalışmanın odağı faz dedektörü tabanlı iki kanallı Cross-PSD "
        "yaklaşımıdır.",
    )


def build_model_and_development(document: Document) -> None:
    add_section_title(
        document, "4. İki kanallı MATLAB benzetim modeli", page_break=False
    )
    add_subheading(document, "4.1. Gürültü ve taşıyıcı üretimi")
    add_body_paragraph(
        document,
        "Her Monte Carlo yinelemesinde yeni bir DUT faz dizisi ile iki ayrı "
        "referans faz dizisi üretilir. Aynı DUT taşıyıcısı iki ölçüm kanalına "
        "uygulanır; referansların merkez fazına π/2 eklenir. Ayrı randn "
        "örnekleri kullanılmakla birlikte bağımsızlık için otomatik kabul "
        "testi bulunmadığından referanslar raporda 'modelde korelasyonsuz kabul "
        "edilen diziler' olarak tanımlanmıştır.",
    )
    add_process_chain(document)

    add_subheading(document, "4.2. Mikser, LPF ve faz dedektörü kazancı")
    add_body_paragraph(
        document,
        "DUT her referansla örnek örnek çarpılır. Mikser çıkışı faz farkını "
        "taşıyan taban bant bileşeniyle yaklaşık 2f₀ çevresindeki toplam "
        "frekans bileşenini birlikte içerir. Dördüncü derece nedensel "
        "Butterworth filter() çağrısı toplam bileşeni bastırır; filtfilt "
        "kullanılmaz. Çıkış küçük-sinyal faz dedektörü kazancına bölünür:",
    )
    add_equation(document, "Kpd = A² / 2")
    add_body_paragraph(
        document,
        "Filtre sonrası seçilen başlangıç örnekleri atılır ve iki kanalın DC "
        "bileşenleri ayrı ayrı çıkarılır. Final profillerde settling_samples=0 "
        "olduğundan IIR başlangıç geçicisinin ayrıca elenmediği sınırlamalar "
        "bölümünde belirtilmiştir.",
    )

    add_subheading(document, "4.3. Kompleks Cross-PSD ve doğrusal ortalama")
    add_body_paragraph(
        document,
        "İki kanalın FFT'si alınır ve tek taraflı kompleks çapraz spektrum "
        "hesaplanır. DC ve Nyquist kutuları tek taraflı dönüşümde ikiyle "
        "çarpılmaz.",
    )
    add_equation(document, "Ŝ₁₂[k] = X₁[k] X₂*[k] / (fs M)")
    add_body_paragraph(
        document,
        "Kompleks spektrumlar bütün yinelemelerde toplanır; büyüklük işlemi "
        "kompleks ortalamadan sonra uygulanır. Aynı yinelemelerdeki filtresiz "
        "DUT faz dizilerinin dikdörtgen pencereli periodogramları da doğrusal "
        "güç alanında ortalanır. Böylece iki eğri aynı Monte Carlo "
        "popülasyonuna dayanır.",
    )

    add_subheading(document, "4.4. Logaritmik binleme ve değerlendirme metriği")
    add_body_paragraph(
        document,
        "Bin merkezleri geometrik frekans ortalamasıyla, bin gücü ise doğrusal "
        "PSD'nin aritmetik ortalamasıyla elde edilir. SSB dBc/Hz dönüşümü "
        "binlemeden sonra uygulanır. MAE, iki log-bin eğrisinin ortak "
        "frekans aralığında 200 logaritmik noktaya enterpolasyonundan sonra "
        "hesaplanan ortalama mutlak dB farkıdır:",
    )
    add_equation(document, "MAE = (1/J) Σⱼ |Lcross(fⱼ) − LDUT(fⱼ)|")
    add_callout(
        document,
        "Metrik sınırı",
        "Cross-PSD kanalları filtreli, DUT karşılaştırma periodogramı "
        "filtresizdir. Resmî MAE bütün ortak pozitif frekans bandını kapsar; "
        "LPF geçiş/durdurma bandı ve mikser toplam-frekans kalıntısı da metriğe "
        "katılabilir. Bu nedenle MAE saf korelasyon hatası veya donanım "
        "doğruluğu değildir.",
        color=LIGHT_BLUE,
    )

    add_subheading(document, "4.5. Model kapsamı")
    add_body_paragraph(
        document,
        "R&S donanım zincirindeki PLL, LNA ve ADC blokları güncel MATLAB "
        "modelinde ayrıca gerçekleştirilmemiştir. Kuantalama, örnekleme saati "
        "jitteri, kanal sızıntısı, kazanç/faz uyumsuzluğu, sıcaklık ve gerçek "
        "mikser doğrusal-olmama etkileri ideal kabul edilmiştir. Sonuçlar ticari "
        "bir analizörün mutlak duyarlılığı yerine yöntem davranışını temsil "
        "eder.",
    )

    add_section_title(
        document,
        "5. Projenin gelişimi ve algoritmik optimizasyonlar",
        page_break=False,
    )
    add_subheading(document, "5.1. Kuramdan ilk prototipe")
    add_body_paragraph(
        document,
        "Çalışma, yerel osilatör faz gürültüsünün sistem performansına etkisi "
        "ve düşük seviyeli ölçümlerde analizör tabanının oluşturduğu sınırın "
        "incelenmesiyle başlamıştır. R&S ve benzeri kaynaklar üzerinden SSB "
        "gösterimi, flicker mekanizmaları, faz dedektörü ve iki kanallı ölçüm "
        "ilkesi çalışılmıştır. Ardından beyaz Gauss gürültüsünün taşıyıcının "
        "fazına eklenmesi, renkli gürültü üretimi ve proje için seçilen 1/f³ "
        "faz PSD modeline geçilmiştir.",
    )
    add_subheading(document, "5.2. Modüler iki kanallı yapı")
    add_body_paragraph(
        document,
        "İlk faz gürültüsü üreticisinin ardından DUT, Ref1 ve Ref2 için ayrı "
        "gerçekleşimler oluşturulmuş; mikser ve LPF blokları eklenmiş; iki kanal "
        "Cross-PSD ile birleştirilmiştir. Kod daha sonra "
        "run_simulation(config) çevresinde modülerleştirilmiş, her yinelemede "
        "yeni DUT üretimine geçilmiş ve DUT periodogramları doğrusal güç "
        "alanında ortalanmıştır.",
    )
    add_subheading(document, "5.3. Hesaplama zincirinin sadeleştirilmesi")
    add_body_paragraph(
        document,
        "İlk yaklaşımda xcorr ile açık korelasyon dizisi oluşturuluyor, "
        "ifftshift uygulanıyor ve yeniden FFT alınıyordu. Güncel yaklaşım "
        "korelasyon teoremine dayalı X₁·X₂* çarpımını doğrudan hesaplar. Aynı "
        "deterministik veriyle otomatik eşdeğerlik testi bulunmadığı için "
        "'tamamen eşdeğer' ifadesi kullanılmamıştır.",
    )
    add_equation(document, "nfft = 2^nextpow2(2M − 1)")
    add_body_paragraph(
        document,
        "Sıfır doldurma frekans ızgarasını sıklaştırır; kayıt süresinin "
        "belirlediği bağımsız fiziksel çözünürlüğü artırmaz. LPF katsayıları "
        "config değişmediği sürece önbellekte tutulmuş, iki kanal tek matris "
        "olarak filtrelenmiş, sabit hesaplar döngü dışına çıkarılmış, log-bin "
        "içinde tepe yerine doğrusal ortalama güç kullanılmış ve aktif olmayan "
        "Welch çıktıları kaldırılmıştır.",
    )
    add_data_table(
        document,
        ["İlk yaklaşım", "Güncel yaklaşım", "Mühendislik etkisi"],
        [
            ["xcorr → ifftshift → FFT", "X₁·X₂* ile doğrudan Cross-PSD", "Açık korelasyon dizisi kaldırıldı"],
            ["nfft = 2M−1", "2^nextpow2(2M−1)", "Radix-2 hesap ve sık frekans ızgarası"],
            ["LPF tasarımı her yinelemede", "Katsayılar config'e göre önbellekte", "Tekrarlı tasarım kaldırıldı"],
            ["Kanallar ayrı filtreleniyor", "İki sütun tek filter() çağrısında", "Çağrı yükü azaltıldı"],
            ["Ara adımda bölerek ortalama", "Topla, döngü sonunda böl", "Döngü içi bölmeler azaltıldı"],
            ["Log-bin içinde tepe değer", "Doğrusal PSD aritmetik ortalaması", "Güç ortalaması korunuyor"],
            ["Kullanılmayan Welch çıktıları", "Cross-PSD ve DUT periodogramı", "Çıktı zinciri sadeleşti"],
            ["Kontrolsüz config", "validate_config doğrulaması", "Hatalar büyük dizilerden önce yakalanıyor"],
        ],
        [5.1, 6.2, 4.7],
    )
    add_subheading(document, "5.4. MATLAB R2025b uyarlaması ve uzun koşular")
    add_body_paragraph(
        document,
        "Commitli optimized kaynak GNU Octave uygulamasıdır. Rapor sonuçları, "
        f"{SOURCE_COMMIT} senkron noktasından türetilen matlab_version "
        "klasöründe MATLAB R2025b ile üretilmiştir. Octave'a özgü pkg load "
        "signal, time() ve -mat7-binary kullanımları kaldırılmış; MATLAB random "
        "stream'i oturumda bir kez rng(\"shuffle\") ile başlatılmış; büyük "
        "sonuçlar -v7.3 ve grafikler 150 DPI PNG olarak kaydedilmiştir.",
    )
    add_body_paragraph(
        document,
        "Uzun yineleme taramalarında tamamlanan noktaları korumak için config "
        "uyumluluğunu denetleyen bir sonuç birleştirme akışı geliştirilmiştir. "
        "Kaynak klasörler değiştirilmeden yalnız eksik noktaların çalıştırılması "
        "hedeflenmiştir. Rapor oluşturucu da final uzun koşu tamamlanmadan veya "
        "N=1.000.000 config'i doğrulanmadan belge üretmez.",
    )
    add_subheading(document, "5.5. Commitlere dayalı değişiklik özeti")
    add_data_table(
        document,
        ["Commit", "Tarih", "Doğrulanmış değişiklik", "Rapordaki karşılığı"],
        [
            ["0fa4d8f", "04.08.2026", "AWGN/renkli gürültü ve ilk model", "İlk prototip"],
            ["045fee4", "17.08.2026", "Optimize alan, config ve sweep altyapısı", "Deney otomasyonu"],
            ["b979435", "18.08.2026", "run_simulation API ve çıktı sözleşmesi", "Modüler mimari"],
            ["c8c5260", "20.08.2026", "Yeni DUT/iterasyon, ayrı mikser-LPF-Cross-PSD", "Adil Monte Carlo"],
            ["9b14de0", "20.08.2026", "README, handoff ve kanıt sınırları", "Provenance disiplini"],
            ["ed60dc6", "20.08.2026", "İlk DOCX rapor", "Rapor başlangıcı"],
            ["0799f9f", "21.08.2026", "N=1M profil, merge akışı ve rapor güncellemesi", "Final deney temeli"],
        ],
        [2.2, 2.3, 7.1, 4.6],
        numeric_columns=(0, 1),
    )
    add_callout(
        document,
        "Performans iddiası",
        "Kod yapısı işlem zincirini sadeleştirir; ancak güncel repoda aynı "
        "deterministik veri ve kontrollü sistem yüküyle yapılmış bir benchmark "
        "bulunmadığından hızlanma çarpanı verilmemiştir.",
        color=LIGHT_BLUE,
    )


def build_experiment_design(
    document: Document, sweeps: Mapping[str, SweepData], final_iterations: SweepData
) -> None:
    add_section_title(document, "6. Deney tasarımı ve sonuç sözleşmesi", page_break=True)
    add_subheading(document, "6.1. Genel parametre karşılaştırma profili")
    add_body_paragraph(
        document,
        "run_comparisons.m beş bağımsız tek-parametre taraması yürütür. "
        "Tarama noktaları Kartezyen çarpım değildir; her koşul temiz bir temel "
        "config kopyasından başlar ve yalnız ilgili alan değiştirilir.",
    )
    add_data_table(
        document,
        ["Parametre", "Değer"],
        [
            ["Örnek sayısı N", "1.000.000"],
            ["Örnekleme / taşıyıcı", "1 MHz / 200 kHz"],
            ["Taşıyıcı genliği", "1"],
            ["Atılan geçici rejim", "0 örnek"],
            ["LPF", "4. derece Butterworth, varsayılan 200 kHz"],
            ["DUT RMS", "0,05 rad"],
            ["Ref1 / Ref2 RMS", "0,05 / 0,05 rad"],
            ["Temel yineleme", "100"],
            ["Log-bin", "100"],
        ],
        [6.4, 9.6],
    )
    add_data_table(
        document,
        ["Tarama", "Değerler"],
        [
            ["LPF kesimi", "1, 5, 10, 25, 50, 75, 100, 200, 300 kHz"],
            ["DUT RMS", "0,01; 0,02; 0,05; 0,10; 0,20; 0,50 rad"],
            ["Ref1=Ref2 RMS", "0,01; 0,02; 0,05; 0,10; 0,20; 0,50 rad"],
            ["Yineleme", "1, 10, 100, 200, 500, 1000"],
            ["Log-bin", "10, 25, 50, 80, 100, 200"],
        ],
        [5.2, 10.8],
    )

    add_subheading(document, "6.2. Özel uzun yineleme profili")
    add_body_paragraph(
        document,
        "Uzun tarama, genel karşılaştırma profilinden farklı bir duyarlılık "
        "senaryosudur. İki seri aynı kontrollü eğri gibi birleştirilmemiştir. "
        "Rapor oluşturucu final klasörünü çalışma anında keşfeder ve her raw "
        "MAT config'ini aşağıdaki sözleşmeye karşı doğrular.",
    )
    add_data_table(
        document,
        ["Parametre", "Değer"],
        [
            ["Sonuç klasörü", final_iterations.directory.name],
            ["Örnek sayısı N", "1.000.000"],
            ["Örnekleme / taşıyıcı", "1 MHz / 200 kHz"],
            ["LPF", "4. derece Butterworth, 100 kHz"],
            ["DUT RMS", "0,02 rad"],
            ["Ref1 / Ref2 RMS", "0,05 / 0,05 rad"],
            ["Log-bin", "100"],
            ["Yineleme değerleri", ", ".join(fmt_int(v) for v in EXPECTED_FINAL_ITERATIONS)],
        ],
        [6.4, 9.6],
    )
    add_body_paragraph(
        document,
        "DUT ve referanslar aynı 1/f³ spektral biçimine göre ölçeklendiği için "
        "bir referansın teorik PSD seviyesi DUT seviyesinden yaklaşık 7,96 dB "
        "yüksektir:",
    )
    add_equation(
        document,
        "ΔL = 20 log₁₀(σref/σDUT) = 20 log₁₀(0,05/0,02) ≈ 7,96 dB",
    )

    add_subheading(document, "6.3. Sonuç ve kanıt sözleşmesi")
    add_numbered(document, 1, "summary.csv satırları beklenen değer listesiyle eşleştirilir.")
    add_numbered(document, 2, "Her satırdaki raw MAT dosyası açılır ve N=1.000.000 dahil bütün config alanları doğrulanır.")
    add_numbered(document, 3, "CSV'deki değer, MAE, düzeltme katsayısı ve süre ham MAT içeriğiyle yuvarlama hassasiyetinde karşılaştırılır.")
    add_numbered(document, 4, "summary.mat ve karşılaştırma PNG'sinin varlığı; PNG'nin açılabilirliği ve asgari çözünürlüğü denetlenir.")
    add_numbered(document, 5, "Final uzun tarama beklenen dokuz noktayı içermiyorsa DOCX yazılmaz.")
    add_callout(
        document,
        "Veri ayrımı",
        f"Karşılaştırma batch'i {COMPARISON_STAMP}_*; final uzun yineleme "
        f"klasörü {final_iterations.directory.name} olarak doğrulanmıştır. "
        "İki profilin sabit LPF, DUT RMS ve temel yineleme değerleri farklıdır.",
    )


def build_results(
    document: Document, sweeps: Mapping[str, SweepData], final_iterations: SweepData
) -> None:
    add_section_title(document, "7. N=1.000.000 benzetim sonuçları", page_break=True)
    comparison_total = sum(sweep.total_elapsed_s for sweep in sweeps.values())
    add_subheading(document, "7.1. Veri bütünlüğü")
    add_body_paragraph(
        document,
        "Beş karşılaştırma taramasındaki bütün CSV satırları ilgili raw MATLAB "
        "v7.3 dosyalarıyla ve config alanlarıyla doğrulanmıştır. Beş "
        "karşılaştırma grafiği ile final uzun-yineleme grafiği PNG olarak "
        "açılmıştır. Karşılaştırma batch'indeki kayıtlı simülasyon sürelerinin "
        f"toplamı {fmt_duration(comparison_total)}, final uzun taramanın "
        f"kayıtlı toplam süresi {fmt_duration(final_iterations.total_elapsed_s)} "
        "olmuştur. Dosya yazma ve grafik üretme ek süreleri bu toplamların "
        "dışında olabilir.",
    )

    # LPF
    lpf = sweeps["lpf_cutoff"]
    lpf_best = best_row(lpf)
    lpf_worst = worst_row(lpf)
    add_subheading(document, "7.2. LPF kesim frekansının etkisi")
    add_data_table(
        document,
        ["LPF (kHz)", "MAE (dB)", "Düzeltme", "Süre (s)"],
        result_table_rows(lpf, lambda value: fmt_int(value / 1_000)),
        [3.4, 3.6, 4.2, 4.2],
        numeric_columns=(0, 1, 2, 3),
    )
    add_body_paragraph(
        document,
        f"Bu taramada en düşük gözlenen tam bant MAE "
        f"{fmt_int(lpf_best.value / 1_000)} kHz'de "
        f"{fmt_decimal(lpf_best.mae_db)} dB, en yüksek değer "
        f"{fmt_int(lpf_worst.value / 1_000)} kHz'de "
        f"{fmt_decimal(lpf_worst.mae_db)} dB'dir. Tek rastgele koşuya ve "
        "filtresiz DUT karşılaştırmasına dayandığı için en düşük nokta genel "
        "bir filtre optimumu olarak yorumlanmamıştır.",
    )
    add_body_paragraph(
        document,
        "300 kHz kesiminde 2f₀=400 kHz toplam-frekans bileşeninin dördüncü "
        "derece LPF tarafından daha az bastırılması gözlenen geniş bant "
        "ayrışmayla uyumlu olabilir. Ayrı bir filtre tepki ölçümü yapılmadığı "
        "için bu yorum nitel düzeyde tutulmuştur.",
    )
    add_figure(
        document,
        lpf.plot_path,
        "Şekil 7.1. Dokuz LPF kesim frekansında Cross-PSD ve aynı "
        "yinelemelerin filtresiz DUT periodogramı. N=1.000.000; temel "
        "DUT/Ref RMS=0,05 rad; 100 yineleme. Kaynak: Bu çalışma.",
        max_height_cm=13.2,
    )

    # DUT RMS
    dut = sweeps["rms_dut"]
    dut_best = best_row(dut)
    dut_high = get_row(dut, 0.50)
    add_subheading(document, "7.3. DUT RMS değerinin etkisi")
    add_data_table(
        document,
        ["DUT RMS (rad)", "MAE (dB)", "Düzeltme", "Süre (s)"],
        result_table_rows(dut, lambda value: fmt_decimal(value, 2)),
        [3.7, 3.5, 4.2, 4.0],
        numeric_columns=(0, 1, 2, 3),
    )
    add_body_paragraph(
        document,
        f"Sabit 0,05/0,05 rad referans koşulunda en düşük gözlenen MAE "
        f"{fmt_decimal(dut_best.value, 2)} rad DUT RMS noktasında "
        f"{fmt_decimal(dut_best.mae_db)} dB'dir. 0,50 rad noktasında MAE "
        f"{fmt_decimal(dut_high.mae_db)} dB'ye yükselirken düzeltme katsayısı "
        f"{fmt_decimal(dut_high.correction_factor, 6)} olmuştur. Katsayının "
        f"birden sapması yaklaşık %{fmt_decimal((dut_high.correction_factor - 1) * 100, 1)} "
        "olup yüksek faz RMS değerlerinde küçük-açı yaklaşımından uzaklaşmanın "
        "önem kazandığını gösterir. Bu nokta genel donanım çalışma sınırı "
        "değildir.",
    )
    add_figure(
        document,
        dut.plot_path,
        "Şekil 7.2. Farklı DUT RMS değerlerinde Cross-PSD ve filtresiz DUT "
        "periodogramı. N=1.000.000; Ref1=Ref2=0,05 rad; 100 yineleme. "
        "Kaynak: Bu çalışma.",
        max_height_cm=9.2,
    )

    # Reference RMS
    reference = sweeps["rms_ref"]
    reference_default = get_row(reference, 0.05)
    reference_high = get_row(reference, 0.50)
    reference_steps = increasing_steps(reference.rows)
    add_subheading(document, "7.4. Referans RMS değerinin etkisi")
    add_data_table(
        document,
        ["Ref1=Ref2 RMS (rad)", "MAE (dB)", "Düzeltme", "Süre (s)"],
        result_table_rows(reference, lambda value: fmt_decimal(value, 2)),
        [4.2, 3.4, 4.2, 3.6],
        numeric_columns=(0, 1, 2, 3),
    )
    add_body_paragraph(
        document,
        f"Referans RMS 0,05 rad iken MAE {fmt_decimal(reference_default.mae_db)} "
        f"dB, 0,50 rad iken {fmt_decimal(reference_high.mae_db)} dB'dir; "
        f"mutlak fark {fmt_decimal(reference_high.mae_db - reference_default.mae_db)} "
        "dB'dir. Bu taramada yüksek referans gürültüsünün sonlu 100 "
        "yinelemede daha fazla artık bileşen oluşturduğu görülmüştür. "
        f"Değer sırası incelendiğinde {describe_steps(reference_steps)}. "
        "Sonuç referansların istatistiksel bağımsızlığını tek başına "
        "kanıtlamaz.",
    )
    add_figure(
        document,
        reference.plot_path,
        "Şekil 7.3. İki referans RMS değerinin birlikte taranması. "
        "N=1.000.000; DUT RMS=0,05 rad; 100 yineleme. Kaynak: Bu çalışma.",
        max_height_cm=9.2,
    )

    # General comparison iterations
    comparison_iterations = sweeps["iterations"]
    ci_first = comparison_iterations.rows[0]
    ci_last = comparison_iterations.rows[-1]
    ci_steps = increasing_steps(comparison_iterations.rows)
    add_subheading(document, "7.5. Genel profilde yineleme sayısının etkisi")
    add_data_table(
        document,
        ["Yineleme", "MAE (dB)", "Düzeltme", "Süre (s)"],
        result_table_rows(comparison_iterations, fmt_int),
        [3.4, 3.6, 4.2, 4.2],
        numeric_columns=(0, 1, 2, 3),
    )
    add_body_paragraph(
        document,
        f"Yineleme sayısı {fmt_int(ci_first.value)}'den "
        f"{fmt_int(ci_last.value)}'e çıkarıldığında tam bant MAE "
        f"{fmt_decimal(ci_first.mae_db)} dB'den {fmt_decimal(ci_last.mae_db)} "
        f"dB'ye düşmüş; {describe_steps(ci_steps)}. Genel eğilim ortak DUT "
        "spektrumuna yaklaşmayı desteklese de her artışın daha düşük MAE "
        "üreteceği iddia edilmemiştir.",
    )
    add_figure(
        document,
        comparison_iterations.plot_path,
        "Şekil 7.4. Genel karşılaştırma profilinde 1–1.000 yineleme. "
        "N=1.000.000; LPF=200 kHz; DUT/Ref RMS=0,05 rad. "
        "Kaynak: Bu çalışma.",
        max_height_cm=9.2,
    )

    # Log bins
    log_bins = sweeps["log_bins"]
    bins_best = best_row(log_bins)
    bins_steps = increasing_steps(log_bins.rows)
    add_subheading(document, "7.6. Logaritmik bin sayısının etkisi")
    add_data_table(
        document,
        ["Log-bin", "MAE (dB)", "Düzeltme", "Süre (s)"],
        result_table_rows(log_bins, fmt_int),
        [3.4, 3.6, 4.2, 4.2],
        numeric_columns=(0, 1, 2, 3),
    )
    add_body_paragraph(
        document,
        f"Bu tek taramada en düşük gözlenen MAE {fmt_int(bins_best.value)} "
        f"bin için {fmt_decimal(bins_best.mae_db)} dB'dir ve "
        f"{describe_steps(bins_steps)}. Her bin değeri için yeni rastgele "
        "gerçekleşimler üretildiğinden fark yalnız binleme çözünürlüğüne "
        "bağlanamaz. Aynı tam çözünürlüklü PSD'nin yeniden binlenmesi "
        "yapılmadan istatistiksel optimum iddiası kurulmamıştır.",
    )
    add_figure(
        document,
        log_bins.plot_path,
        "Şekil 7.5. Logaritmik bin sayısının spektral gösterime ve tam bant "
        "MAE'ye etkisi. N=1.000.000; temel profil. Kaynak: Bu çalışma.",
        max_height_cm=9.2,
    )

    # Final long iterations
    long_first = final_iterations.rows[0]
    long_last = final_iterations.rows[-1]
    long_steps = increasing_steps(final_iterations.rows)
    row_10k = get_row(final_iterations, 10_000.0)
    row_20k = get_row(final_iterations, 20_000.0)
    low_band_10k = compute_band_mae(
        final_iterations.directory / "raw" / row_10k.run_file, 10_000.0
    )
    low_band_20k = compute_band_mae(
        final_iterations.directory / "raw" / row_20k.run_file, 10_000.0
    )
    add_subheading(document, "7.7. Özel uzun yineleme taraması")
    add_data_table(
        document,
        ["Yineleme", "MAE (dB)", "Düzeltme", "Süre (s)"],
        result_table_rows(final_iterations, fmt_int),
        [3.4, 3.6, 4.2, 4.2],
        numeric_columns=(0, 1, 2, 3),
    )
    long_change = long_first.mae_db - long_last.mae_db
    last_delta = row_20k.mae_db - row_10k.mae_db
    add_body_paragraph(
        document,
        f"Referans PSD seviyesinin DUT'tan teorik olarak yaklaşık 7,96 dB "
        f"yüksek olduğu özel profilde, {fmt_int(long_first.value)} yinelemede "
        f"{fmt_decimal(long_first.mae_db)} dB olan tam bant MAE "
        f"{fmt_int(long_last.value)} yinelemede {fmt_decimal(long_last.mae_db)} "
        f"dB olmuştur. İlk-son azalma {fmt_decimal(long_change)} dB'dir ve "
        f"{describe_steps(long_steps)}.",
    )
    add_body_paragraph(
        document,
        f"10.000→20.000 geçişindeki MAE değişimi "
        f"{fmt_decimal(last_delta, 3)} "
        f"dB, gözlenen süre oranı {fmt_decimal(row_20k.elapsed_s / row_10k.elapsed_s, 3)} "
        "olmuştur. MAE değişiminin işareti tek başına 20.000 noktasının daha "
        "iyi veya kötü olduğunu kanıtlamaz; final değerlendirme tek koşulu "
        "Monte Carlo dalgalanması ve geniş bant sistematik tabanla birlikte "
        "yapılmıştır.",
    )
    add_callout(
        document,
        "Düşük-offset yardımcı kontrol",
        f"Resmî tam-bant metriğe ek olarak, bu metrik için kurulan 200 "
        f"noktalı ortak log-frekans ızgarasının 10 kHz ve altındaki alt "
        f"kümesinde MAE; 10.000 yinelemede "
        f"{fmt_decimal(low_band_10k)} dB, 20.000 yinelemede "
        f"{fmt_decimal(low_band_20k)} dB bulunmuştur. Bu değer LPF üstü "
        "beklenen ayrışmayı dışlayan ek bir tanıdır; resmî metrik yerine "
        "geçmez.",
        color=LIGHT_BLUE,
    )

    # Use a landscape page so the multi-panel iteration plot remains legible.
    start_landscape_section(document)
    heading = document.add_heading("7.7.1. Uzun tarama karşılaştırma grafiği", level=2)
    add_figure(
        document,
        final_iterations.plot_path,
        "Şekil 7.6. Özel N=1.000.000 profilde final uzun yineleme taraması. "
        "LPF=100 kHz; DUT RMS=0,02 rad; Ref1=Ref2=0,05 rad; 100 log-bin. "
        "Kaynak: Bu çalışma.",
        max_width_cm=25.2,
        max_height_cm=13.8,
    )

    start_portrait_section(document)


def build_discussion_conclusion_and_appendices(
    document: Document,
    sweeps: Mapping[str, SweepData],
    final_iterations: SweepData,
) -> None:
    # A portrait section has already been created after the landscape result page.
    heading = document.add_heading("8. Tartışma ve sınırlamalar", level=1)
    set_paragraph_border(heading, "bottom", TURQUOISE, size="10")
    final_first = final_iterations.rows[0]
    final_last = final_iterations.rows[-1]
    final_steps = increasing_steps(final_iterations.rows)
    final_low_band_mae = compute_band_mae(
        final_iterations.directory / "raw" / final_last.run_file, 10_000.0
    )
    ref_high = get_row(sweeps["rms_ref"], 0.50)
    ref_default = get_row(sweeps["rms_ref"], 0.05)
    dut_high = get_row(sweeps["rms_dut"], 0.50)

    add_subheading(document, "8.1. Kuramsal beklentiyle uyum")
    add_body_paragraph(
        document,
        f"Özel uzun taramada tam bant MAE {fmt_decimal(final_first.mae_db)} "
        f"dB'den {fmt_decimal(final_last.mae_db)} dB'ye düşmüş ve "
        f"{describe_steps(final_steps)}. Genel azalma, kanala özgü referans "
        "terimlerinin kompleks ortalamayla bastırılması yönündeki kuramsal "
        "beklentiyle uyumludur. Ancak R&S'nin 5 log₁₀(K) bağıntısı gürültü "
        "tabanı iyileşmesini ifade eder; bu çalışmanın filtresiz DUT "
        "karşılaştırmasını da içeren geniş bant MAE'sine doğrudan "
        "uygulanmamıştır.",
    )
    add_body_paragraph(
        document,
        f"20.000 yineleme sonucunda 10 kHz ve altı için hesaplanan ek "
        f"tanısal MAE {fmt_decimal(final_low_band_mae)} dB'dir. Bu belirgin "
        f"düşük-offset uyumu, tam bant MAE'nin {fmt_decimal(final_last.mae_db)} "
        "dB düzeyinde kalmasında LPF üstü ayrışma ve 2f₀ kalıntısının etkili "
        "olduğu yorumuyla uyumludur; "
        "ayrı bir belirsizlik analizi yerine geçmez.",
    )

    add_subheading(document, "8.2. LPF ve mikser toplam-frekans bileşeni")
    add_body_paragraph(
        document,
        "LPF taraması, ölçüm kanalının bant sınırlamasının metrik üzerinde "
        "belirleyici olduğunu göstermiştir. Cross-PSD eğrisi filtreli, DUT "
        "periodogramı filtresiz olduğu için kesim üzerindeki ayrışma kısmen "
        "beklenir. 2f₀ çevresindeki dar kalıntı yüksek kesimlerde görünür hale "
        "gelebilir; bu bileşen DUT faz gürültüsü veya korelasyon tabanı olarak "
        "yorumlanmamalıdır.",
    )

    add_subheading(document, "8.3. DUT ve referans seviye dengesi")
    add_body_paragraph(
        document,
        f"Referans RMS'in 0,05 rad'dan 0,50 rad'a çıkarılmasıyla MAE "
        f"{fmt_decimal(ref_default.mae_db)} dB'den {fmt_decimal(ref_high.mae_db)} "
        "dB'ye yükselmiştir. Sonlu ortalamada yüksek referans gürültüsü daha "
        "fazla yineleme gerektirir. DUT RMS taramasında yüksek ortak bileşen "
        "başlangıçta uyumu iyileştirmiş, fakat 0,50 rad'da düzeltme katsayısı "
        f"{fmt_decimal(dut_high.correction_factor, 6)} değerine ulaşmıştır. "
        "Bu davranış küçük-açı yaklaşımının sınırsız geçerli olmadığını "
        "gösterir.",
    )

    add_subheading(document, "8.4. İstatistiksel ve sayısal sınırlamalar")
    add_bullet(
        document,
        "Her tarama noktası tek rastgele koşudan oluşur; güven aralığı ve "
        "çoklu tekrar yoktur.",
    )
    add_bullet(
        document,
        "rng(\"shuffle\") başlangıç durumu sonuç dosyalarına kaydedilmez; "
        "koşular bit düzeyinde yeniden üretilemez.",
    )
    add_bullet(
        document,
        "Ref1 ve Ref2 ayrı dizilerden üretilir; istatistiksel bağımsızlık için "
        "otomatik korelasyon kabul testi yoktur.",
    )
    add_bullet(
        document,
        "settling_samples=0 olduğu için IIR başlangıç geçicisi ayrıca atılmaz.",
    )
    add_bullet(
        document,
        "Log-bin taraması her nokta için yeni veriler üretir; yalnız binleme "
        "etkisini izole etmez.",
    )
    add_bullet(
        document,
        "Sıfır doldurma frekans ızgarasını sıklaştırır, bağımsız fiziksel "
        "çözünürlüğü artırmaz.",
    )

    add_subheading(document, "8.5. Donanım soyutlamaları")
    add_body_paragraph(
        document,
        "Model PLL, LNA, ADC kuantalaması, örnekleme saati jitteri, kanal "
        "kazanç/faz uyumsuzluğu, sızıntı, sıcaklık ve kalibrasyon hatalarını "
        "içermez. Gürültü üreticisi tek bir 1/f³ eğimi ve hedef toplam RMS "
        "kullanır. Bu nedenle sonuçlar gerçek bir analizörün mutlak gürültü "
        "tabanı veya ölçüm belirsizliği olarak sunulamaz.",
    )

    add_section_title(document, "9. Sonuç ve gelecek çalışmalar", page_break=True)
    long_first = final_iterations.rows[0]
    long_last = final_iterations.rows[-1]
    lpf_best = best_row(sweeps["lpf_cutoff"])
    bins_best = best_row(sweeps["log_bins"])
    add_body_paragraph(
        document,
        "Bu projede iki referans kanallı faz dedektörü mimarisi MATLAB R2025b "
        "ortamında gerçekleştirilmiş, DUT faz gürültüsü kompleks Cross-PSD "
        "ortalamasıyla kestirilmiş ve aynı Monte Carlo popülasyonunun filtresiz "
        "DUT periodogramıyla karşılaştırılmıştır. Fiziksel çapraz korelasyon "
        "ilkesi ile kodda kullanılan frekans bölgesi Cross-PSD kestiricisi "
        "ayrı biçimde tanımlanmıştır.",
        style="Lead",
    )
    add_body_paragraph(
        document,
        f"Final N=1.000.000 uzun taramasında yineleme sayısı "
        f"{fmt_int(long_first.value)}'den {fmt_int(long_last.value)}'e "
        f"çıkarıldığında tam bant MAE {fmt_decimal(long_first.mae_db)} dB'den "
        f"{fmt_decimal(long_last.mae_db)} dB'ye düşmüştür. LPF taramasında "
        f"en düşük gözlenen MAE {fmt_int(lpf_best.value / 1_000)} kHz'de "
        f"{fmt_decimal(lpf_best.mae_db)} dB; log-bin taramasında en düşük "
        f"gözlenen değer {fmt_int(bins_best.value)} binde "
        f"{fmt_decimal(bins_best.mae_db)} dB olmuştur. Bu noktalar yalnız tek "
        "taramanın gözlemleridir; genel optimum olarak sunulmamıştır.",
    )
    add_body_paragraph(
        document,
        "Açık korelasyon dizisinin FFT tabanlı spektral çarpımla değiştirilmesi, "
        "radix-2 FFT boyu, LPF katsayı önbelleği, iki kanalın birlikte "
        "filtrelenmesi ve ham/özet/görsel sonuç sözleşmesi uzun deneylerin "
        "yönetimini kolaylaştırmıştır. Kontrollü benchmark yapılmadığı için "
        "hızlanma çarpanı verilmemiştir.",
    )
    add_body_paragraph(
        document,
        "Sonraki adımlar; RNG seed/state kaydı, bağımsız substream'ler ve "
        "otomatik korelasyon testleri, kullanıcı tanımlı offset-band MAE'si, "
        "aynı config için çoklu tekrar ve güven aralıkları ile PLL/LNA/ADC ve "
        "kanal uyumsuzluklarının modele eklenmesidir. Bu iyileştirmeler "
        "benzetim ile fiziksel analizör arasındaki boşluğu azaltacaktır.",
    )

    add_section_title(document, "Kaynakça", page_break=True)
    references = [
        "[1] Rohde & Schwarz, Mastering Phase Noise Measurements, Parts 1-3, "
        "Application Note. https://cdn.rohde-schwarz.com/am/us/campaigns_2/"
        "embedded/Rohde_Schwarz_Phase_Noise_App_Note_Allparts.pdf",
        "[2] MathWorks, fft - Fast Fourier transform, MATLAB R2025b "
        "Documentation. https://www.mathworks.com/help/matlab/ref/fft.html",
        "[3] MathWorks, butter - Butterworth filter design, Signal Processing "
        "Toolbox Documentation. https://www.mathworks.com/help/signal/ref/butter.html",
        "[4] MathWorks, rng - Control random number generator, MATLAB "
        "Documentation. https://www.mathworks.com/help/matlab/ref/rng.html",
        "[5] J. S. Bendat and A. G. Piersol, Random Data: Analysis and "
        "Measurement Procedures, 4th ed., Wiley, 2010.",
        f"[6] Proje Git geçmişi ve optimized kaynaklar, kaynak senkron "
        f"commit'i {SOURCE_COMMIT}, 21 Ağustos 2026.",
    ]
    for reference in references:
        add_body_paragraph(document, reference)

    add_section_title(
        document, "Ek A. Sonuç provenance ve config manifesti", page_break=True
    )
    add_data_table(
        document,
        ["Kanıt", "Doğrulanmış konum / değer"],
        [
            ["Kaynak DOCX", "İki Kanallı Cross.docx (repo kökü)"],
            ["Kaynak senkron commit'i", SOURCE_COMMIT],
            ["MATLAB çalışma ortamı", RUNTIME_LABEL],
            ["Karşılaştırma batch'i", f"results/{COMPARISON_STAMP}_*"],
            ["Final iterations", f"results/{final_iterations.directory.name}"],
            ["Örnek sayısı", "Her iki profilde N=1.000.000"],
            ["Karşılaştırma sabitleri", "LPF=200 kHz; DUT/Ref=0,05 rad; temel K=100"],
            ["Uzun tarama sabitleri", "LPF=100 kHz; DUT=0,02 rad; Ref=0,05 rad"],
            ["Resmî metrik", "200 log-frekans noktasında ortalama mutlak dB farkı"],
        ],
        [5.3, 10.7],
    )
    add_callout(
        document,
        "Yeniden üretilebilirlik notu",
        "Raw MAT dosyaları gerçek config ve tam spektrumları içerir; ancak RNG "
        "başlangıç state'i kaydedilmediği için eğriler bit düzeyinde yeniden "
        "üretilemez.",
        color=LIGHT_BLUE,
    )

    add_section_title(document, "Ek B. Yazılım bileşenleri ve sorumlulukları", page_break=True)
    add_data_table(
        document,
        ["Dosya", "Sorumluluk"],
        [
            ["run_simulation.m", "Ana Monte Carlo akışı, ortalama, düzeltme, binleme ve metrik"],
            ["validate_config.m", "Zorunlu config alanları ve fiziksel/sayısal sınırlar"],
            ["generate_phase_noise.m", "1/f³ şekillendirme ve RMS normalizasyonu"],
            ["measure_iteration.m", "Tek yinelemede iki ölçüm kanalının bağlanması"],
            ["mixer.m", "Ortak DUT ile iki referans taşıyıcının çarpımı"],
            ["lowpass_filter.m", "Butterworth tasarımı, önbellek ve iki kanal filtreleme"],
            ["compute_cross_psd.m", "Kompleks tek taraflı FFT Cross-PSD"],
            ["compute_periodogram.m", "Filtresiz DUT periodogramı"],
            ["logbin_phase_noise.m", "Doğrusal PSD ortalaması ve SSB dBc/Hz dönüşümü"],
            ["run_comparisons_main.m", "Tek-parametre taraması, kayıt, CSV/MAT/PNG üretimi"],
            ["replot_results_main.m", "Kayıtlı sonuçların simülasyonsuz yeniden çizimi"],
            ["build_report.py", "Sonuç doğrulaması ve nihai DOCX üretimi"],
        ],
        [5.1, 10.9],
    )


def build_document(
    sweeps: Mapping[str, SweepData], final_iterations: SweepData
) -> Document:
    document = Document(str(SOURCE_DOCX))
    clear_document_body(document)
    if not document.sections:
        raise ReportBuildError("Kaynak DOCX içinde bölüm (section) bulunamadı.")
    configure_section(document.sections[0], landscape=False)
    configure_styles(document)
    configure_header_footer(document)

    core = document.core_properties
    core.title = "İki Kanallı Cross-PSD Yöntemiyle Faz Gürültüsü Ölçümünün Simülasyonu"
    core.author = "Ömer"
    core.subject = "MATLAB R2025b teknik proje raporu"
    core.keywords = "faz gürültüsü, Cross-PSD, çapraz korelasyon, MATLAB"
    core.comments = (
        f"Kaynak belge korunarak {SOURCE_COMMIT} senkron noktasından üretilmiştir."
    )

    build_cover(document)
    build_front_matter(document, sweeps, final_iterations)
    build_theory_and_methods(document)
    build_model_and_development(document)
    build_experiment_design(document, sweeps, final_iterations)
    build_results(document, sweeps, final_iterations)
    build_discussion_conclusion_and_appendices(document, sweeps, final_iterations)
    return document


def save_atomically(document: Document, source_hash_before: str) -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    temp_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            prefix=".report_build_",
            suffix=".docx",
            dir=OUTPUT_DOCX.parent,
            delete=False,
        ) as temp_file:
            temp_path = Path(temp_file.name)
        document.save(str(temp_path))
        if not temp_path.is_file() or temp_path.stat().st_size < 100_000:
            raise ReportBuildError(
                f"Üretilen geçici DOCX beklenenden küçük veya eksik: {temp_path}"
            )
        if sha256_file(SOURCE_DOCX) != source_hash_before:
            raise ReportBuildError("Kaynak DOCX beklenmedik biçimde değişti; çıktı yazılmadı.")
        os.replace(temp_path, OUTPUT_DOCX)
        temp_path = None
    finally:
        if temp_path is not None and temp_path.exists():
            temp_path.unlink()


def main() -> int:
    try:
        source_hash_before = sha256_file(require_file(SOURCE_DOCX, "kaynak DOCX"))
        sweeps, final_iterations = load_report_inputs()
        document = build_document(sweeps, final_iterations)
        save_atomically(document, source_hash_before)
    except ReportBuildError as exc:
        print(f"RAPOR OLUŞTURULAMADI: {exc}", file=sys.stderr)
        return 2
    except Exception as exc:  # fail clearly without a partial final artifact
        print(
            "RAPOR OLUŞTURULAMADI: Beklenmeyen hata: "
            f"{type(exc).__name__}: {exc}",
            file=sys.stderr,
        )
        return 3

    print(f"Rapor oluşturuldu: {OUTPUT_DOCX}")
    print(f"Kaynak DOCX korunmuştur: {SOURCE_DOCX}")
    print(f"Karşılaştırma batch'i: {COMPARISON_STAMP}_*")
    print(f"Final iterations: {final_iterations.directory.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
